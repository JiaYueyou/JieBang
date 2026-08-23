"""读取 private_resume 下 PDF/DOCX，生成不包含原文的简历评测索引和规则 pseudo-gold。"""
import argparse, hashlib, json, re
from pathlib import Path

SKILLS = """
Python Java Go C++ Rust JavaScript TypeScript SQL MySQL Redis MongoDB PostgreSQL
Docker Kubernetes Git GitHub Linux FastAPI Django Flask Spring Boot Vue Vue3 React
PyTorch TensorFlow LangChain LlamaIndex Kafka Spark Elasticsearch Neo4j RAG NLP LLM
DeepSeek DeepSort YOLOv8 NumPy Pandas Matplotlib Pydantic SQLAlchemy Celery Alembic
FAISS OpenCV MQTT LangGraph React Native Ollama vLLM GPT BERT GLM Qwen LLaMA
Llama-Factory Llama.cpp C MATLAB EDA Excel PyCharm Keil Keil5 STM32 Altium Designer
Office Embedding StateGraph RESTful API Prompt Engineering Prompt Template
Structured Output Tool Calling MCP OCR VLM 多模态问答 知识图谱 知识库导入
模型微调 模型推理 模型部署 模型量化 深度学习 自然语言处理 序列模型
CNN K-means 聚类 主成分分析 PID 控制算法 图像处理 串口调试工具
目标检测 视频帧处理 车辆跟踪 车道计数 速度估算 内容指纹去重 字段标准化
岗位 JD 清洗 递归切分 来源回溯 防幻觉设计
""".split()

def text_of(path):
    if path.suffix.lower() == '.docx':
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            parts.extend(' | '.join(cell.text for cell in row.cells) for row in table.rows)
        for section in doc.sections:
            parts.extend(p.text for p in section.header.paragraphs)
            parts.extend(p.text for p in section.footer.paragraphs)
        return '\n'.join(parts)
    from pypdf import PdfReader
    return '\n'.join(page.extract_text() or '' for page in PdfReader(str(path)).pages)

def find(pattern, text):
    m = re.search(pattern, text, re.I)
    return m.group(0) if m else ''

def main():
    parser = argparse.ArgumentParser(); parser.add_argument('--input', default='evaluation/private_resume'); parser.add_argument('--output', default='evaluation/datasets/resume_gold.json'); args = parser.parse_args()
    root = Path(__file__).resolve().parents[2]; folder = root / args.input; items = []
    for path in sorted(folder.glob('*')):
        if path.suffix.lower() not in ('.pdf', '.docx'): continue
        text = text_of(path)
        lower_text = text.lower()
        fragments = {"api", "jd", "岗位", "清洗", "output", "tool", "prompt", "structured", "calling", "engineering", "designer"}
        generic = {"c", "sql"}
        skills = []
        for skill in SKILLS:
            if skill.lower() in fragments:
                continue
            # Avoid substring false positives (for example C in CNN or API in RESTful API).
            if skill.lower() in generic:
                if not re.search(rf"(?<![A-Za-z0-9+#]){re.escape(skill.lower())}(?![A-Za-z0-9+#])", lower_text):
                    continue
            elif skill.lower() not in lower_text:
                continue
            skills.append(skill)
        skills = [s for s in skills if s.lower() not in fragments]
        item_id = hashlib.sha256(path.name.encode()).hexdigest()[:12]
        item = {
            'id': item_id, 'file_name': path.name, 'format': path.suffix.lower()[1:], 'sha256_name': item_id,
            'expected_name': find(r'姓名[:：]?\s*([^\s，,]{2,8})', text),
            'expected_email': find(r'[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}', text),
            'expected_phone': find(r'1[3-9]\d{9}', text),
            'expected_target_position': find(r'(?:求职意向|目标岗位)[:：]?\s*([^\n，,]{2,30})', text),
            'expected_education': 'master' if '硕士' in text else ('bachelor' if '本科' in text else ('college' if '大专' in text else '')),
            'expected_skills': skills, 'annotation_method': 'rule_auto', 'status': 'pseudo_gold',
            'actual_name': '', 'actual_email': '', 'actual_phone': '', 'actual_target_position': '', 'actual_education': '', 'actual_skills': [],
        }
        items.append(item)
    output = root / args.output; output.parent.mkdir(parents=True, exist_ok=True); output.write_text(json.dumps({'version': 'resume-gold-v1-auto', 'status': 'pseudo_gold_rule_auto', 'warning': '自动规则生成，未作为最终准确率证明', 'items': items}, ensure_ascii=False, indent=2), encoding='utf-8'); print(f'indexed {len(items)} resume files -> {output}')

if __name__ == '__main__': main()
