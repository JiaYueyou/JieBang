"""
简历服务 —— 简历 CRUD、文件上传解析（DOCX/PDF/DOC + LLM 结构化提取）。
"""
import json
import re
import os
import subprocess
import tempfile
from pathlib import Path
from datetime import datetime
from sqlalchemy.ext.asyncio import AsyncSession
from app.core.config import LLM_API_KEY, LLM_BASE_URL, LLM_MODEL, LLM_TIMEOUT_SECONDS
from app.core.exceptions import ResourceNotFoundError, InvalidParameterError
from app.repositories.resume_repository import ResumeRepository

# 上传文件存储根目录
UPLOAD_DIR = Path(__file__).resolve().parent.parent.parent.parent / "uploads" / "resumes"


def _ensure_upload_dir():
    """确保上传目录存在"""
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


# ========== 文本提取工具函数 ==========

def extract_text_docx(file_content: bytes) -> str:
    """从 DOCX 文件中提取文本（段落 + 表格）"""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(file_content))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_text_pdf(file_content: bytes) -> str:
    """从 PDF 文件中提取文本，处理双栏布局"""
    from io import BytesIO
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            width = page.width or 0
            height = page.height or 0
            if width > height * 1.3:
                left = page.within_bbox((0, 0, width / 2, height))
                right = page.within_bbox((width / 2, 0, width, height))
                left_text = left.extract_text()
                right_text = right.extract_text()
                if left_text:
                    parts.append(left_text)
                if right_text:
                    parts.append(right_text)
            else:
                text = page.extract_text()
                if text:
                    parts.append(text)

            for table in page.extract_tables():
                for row in table:
                    row_text = " | ".join(str(cell).strip() if cell else "" for cell in row)
                    if row_text.strip():
                        parts.append(row_text)

    return "\n".join(parts)


def extract_text_doc(file_content: bytes) -> str:
    """将 DOC 文件通过 LibreOffice 转为 PDF 后提取文本"""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_doc:
        tmp_doc.write(file_content)
        doc_path = tmp_doc.name

    output_dir = tempfile.gettempdir()
    expected_pdf = os.path.join(output_dir, Path(doc_path).stem + ".pdf")
    try:
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, doc_path],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0 or not os.path.exists(expected_pdf):
            raise RuntimeError(f"LibreOffice 转换失败: {result.stderr.decode()}")

        with open(expected_pdf, "rb") as f:
            pdf_content = f.read()
        return extract_text_pdf(pdf_content)
    finally:
        for p in (doc_path, expected_pdf):
            try:
                if os.path.exists(p):
                    os.remove(p)
            except Exception:
                pass


def extract_text(file_content: bytes, filename: str) -> str:
    """根据文件扩展名选择对应的文本提取方式"""
    ext = Path(filename).suffix.lower()
    if ext in (".docx", ".wps"):
        return extract_text_docx(file_content)
    elif ext == ".pdf":
        return extract_text_pdf(file_content)
    elif ext == ".doc":
        try:
            return extract_text_doc(file_content)
        except Exception:
            try:
                return extract_text_docx(file_content)
            except Exception:
                raise InvalidParameterError("无法解析该 .doc 文件，请转换为 .docx 后重新上传")
    else:
        raise InvalidParameterError(f"不支持的简历文件格式: {ext}，支持 .pdf / .docx / .doc / .wps")


# ========== LLM 结构化提取 ==========

RESUME_EXTRACTION_PROMPT = """你是一位专业的简历解析专家。请从以下简历文本中提取结构化信息，以 JSON 格式返回。

要求：
1. 只返回 JSON，不要包含任何其他文字或 markdown 标记
2. 所有字段尽量提取，无法识别的字段用空字符串或空数组
3. 技能(skills)字段中 name 为技能名，category 为类别（如：编程语言、框架、数据库、云原生、AI技术等）
4. 工作经历和项目经历中的 skills 和 technologies 字段为字符串数组
5. 教育经历、工作经历中的日期格式保持原文或转为 "YYYY-MM" 格式
6. 如果简历包含自我评价/个人总结，提取到 self_evaluation 字段
7. work_mode 可选值: "fulltime"(全职) / "intern"(实习) / "remote"(远程)，默认为 "fulltime"

返回的 JSON 结构如下：
{
  "personal_info": {
    "name": "姓名",
    "email": "邮箱地址",
    "phone": "手机号",
    "location": "所在城市"
  },
  "job_intent": {
    "desired_position": "期望职位",
    "desired_city": "期望城市",
    "salary_expectation": "期望薪资",
    "work_mode": "fulltime"
  },
  "education": [
    {"school": "学校名称", "degree": "学历(本科/硕士/博士)", "major": "专业", "start_date": "开始日期", "end_date": "结束日期"}
  ],
  "work_experience": [
    {"company": "公司名称", "position": "职位", "start_date": "开始日期", "end_date": "结束日期", "description": "工作描述", "skills": ["使用的技术/工具"]}
  ],
  "projects": [
    {"name": "项目名称", "role": "担任角色", "description": "项目描述", "technologies": ["使用技术"], "highlights": ["项目亮点"]}
  ],
  "skills": [
    {"name": "技能名称", "category": "技能类别"}
  ],
  "self_evaluation": "自我评价全文"
}

简历文本：
---
{raw_text}
---"""


async def llm_extract(raw_text: str) -> dict | None:
    """调用 DeepSeek LLM 将简历文本转为结构化 JSON，失败返回 None"""
    if not LLM_API_KEY or LLM_API_KEY.startswith("your-"):
        return None

    import httpx

    prompt = RESUME_EXTRACTION_PROMPT.replace("{raw_text}", raw_text)
    payload = {
        "model": LLM_MODEL,
        "messages": [
            {"role": "system", "content": "你是一位专业的简历解析专家，只返回 JSON，不包含其他内容。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"},
    }
    headers = {
        "Authorization": f"Bearer {LLM_API_KEY}",
        "Content-Type": "application/json",
    }

    try:
        async with httpx.AsyncClient(timeout=LLM_TIMEOUT_SECONDS) as client:
            resp = await client.post(
                f"{LLM_BASE_URL}/chat/completions",
                json=payload,
                headers=headers,
            )
            resp.raise_for_status()
            data = resp.json()
            content = data["choices"][0]["message"]["content"]
            return json.loads(content)
    except Exception:
        return None


# ========== 正则降级提取 ==========

def regex_extract(raw_text: str) -> dict:
    """正则表达式提取邮箱、手机号等基本信息，作为 LLM 失败时的降级方案"""
    result: dict = {
        "personal_name": "",
        "personal_email": "",
        "personal_phone": "",
        "personal_location": "",
    }

    email_match = re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', raw_text)
    if email_match:
        result["personal_email"] = email_match.group()

    phone_match = re.search(r'1[3-9]\d{9}', raw_text)
    if phone_match:
        result["personal_phone"] = phone_match.group()

    lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
    for line in lines[:5]:
        name_match = re.match(r'^[一-龥]{2,4}$', line)
        if name_match and not re.search(r'(简历|个人|联系|教育|工作|项目|技能|自我|学校|公司|大学|有限)', line):
            result["personal_name"] = line
            break

    return result


# ========== ResumeService ==========

class ResumeService:
    """简历业务逻辑"""

    def __init__(self, db: AsyncSession):
        self.repo = ResumeRepository(db)
        self.db = db

    async def list_resumes(self, user_id: int) -> list[dict]:
        resumes = await self.repo.list_by_user(user_id)
        return [self._to_dict(r) for r in resumes]

    async def get_detail(self, resume_id: int) -> dict:
        resume = await self.repo.get_by_id(resume_id)
        if not resume:
            raise ResourceNotFoundError("简历不存在")
        return self._to_dict(resume)

    async def create(self, user_id: int, data: dict) -> dict:
        create_data = self._build_resume_data(data)
        create_data["name"] = data.get("name", "我的简历")
        resume = await self.repo.create(user_id, create_data)
        await self.db.commit()
        return self._to_dict(resume)

    async def update(self, resume_id: int, data: dict) -> dict:
        resume = await self.repo.get_by_id(resume_id)
        if not resume:
            raise ResourceNotFoundError("简历不存在")

        update_data = self._build_resume_data(data)
        await self.repo.update(resume, **update_data)
        await self.db.commit()
        await self.db.refresh(resume)
        return self._to_dict(resume)

    def _build_resume_data(self, data: dict) -> dict:
        """将前端传来的结构化字段映射为数据库列名"""
        result: dict = {}

        if data.get("personal_info"):
            pi = data["personal_info"]
            result.update({
                "personal_name": pi.get("name", ""),
                "personal_email": pi.get("email", ""),
                "personal_phone": pi.get("phone", ""),
                "personal_location": pi.get("location", ""),
            })
        if data.get("job_intent"):
            ji = data["job_intent"]
            result.update({
                "desired_position": ji.get("desired_position", ""),
                "desired_city": ji.get("desired_city", ""),
                "salary_expectation": ji.get("salary_expectation", ""),
                "work_mode": ji.get("work_mode", "fulltime"),
            })
        if data.get("education") is not None:
            result["education_list"] = data["education"]
        if data.get("work_experience") is not None:
            result["work_experience_list"] = data["work_experience"]
        if data.get("projects") is not None:
            result["project_list"] = data["projects"]
        if data.get("skills") is not None:
            result["skill_list"] = data["skills"]

        for front_key, db_key in (("name", "name"), ("target_position", "target_position"), ("self_evaluation", "self_evaluation")):
            if data.get(front_key) is not None:
                result[db_key] = data[front_key]

        return result

    async def delete(self, resume_id: int):
        """删除简历"""
        resume = await self.repo.get_by_id(resume_id)
        if not resume:
            raise ResourceNotFoundError("简历不存在")
        await self.repo.delete(resume)
        await self.db.commit()

    async def duplicate(self, resume_id: int) -> dict:
        """复制简历"""
        resume = await self.repo.get_by_id(resume_id)
        if not resume:
            raise ResourceNotFoundError("简历不存在")
        new_resume = await self.repo.duplicate(resume)
        await self.db.commit()
        return self._to_dict(new_resume)

    async def parse_upload(self, user_id: int, file_content: bytes, filename: str) -> dict:
        """
        上传并解析简历文件。
        流程: 保存原文件 → 提取文本 → LLM 结构化 → 正则降级 → 保存数据库
        """
        _ensure_upload_dir()

        # Step 1: 保存原始文件到磁盘
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{filename}"
        user_dir = UPLOAD_DIR / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)
        file_path = user_dir / safe_filename
        file_path.write_bytes(file_content)

        # Step 2: 提取文本
        try:
            raw_text = extract_text(file_content, filename)
        except InvalidParameterError:
            raise
        except Exception as e:
            resume = await self.repo.create(user_id, {
                "name": Path(filename).stem,
                "source_file": filename,
                "source_file_path": str(file_path.relative_to(UPLOAD_DIR.parent)),
            })
            await self.db.commit()
            return {
                "resume": self._to_dict(resume),
                "extracted_skills": [],
                "parse_accuracy": 0.0,
                "raw_text": None,
                "parse_error": f"文件解析失败: {str(e)}",
            }

        # Step 3: LLM 结构化提取
        parsed = await llm_extract(raw_text)
        parse_accuracy = 0.85 if parsed else 0.0
        extracted_skills: list[str] = []

        # Step 4: 构建数据库写入数据
        db_data: dict = {
            "name": Path(filename).stem,
            "source_file": filename,
            "source_file_path": str(file_path.relative_to(UPLOAD_DIR.parent)),
            "raw_text": raw_text,
        }

        if parsed:
            pi = parsed.get("personal_info", {})
            if pi:
                db_data.update({
                    "personal_name": str(pi.get("name", "") or ""),
                    "personal_email": str(pi.get("email", "") or ""),
                    "personal_phone": str(pi.get("phone", "") or ""),
                    "personal_location": str(pi.get("location", "") or ""),
                })

            ji = parsed.get("job_intent", {})
            if ji:
                db_data.update({
                    "desired_position": str(ji.get("desired_position", "") or ""),
                    "desired_city": str(ji.get("desired_city", "") or ""),
                    "salary_expectation": str(ji.get("salary_expectation", "") or ""),
                    "work_mode": str(ji.get("work_mode", "fulltime") or "fulltime"),
                })

            if parsed.get("education"):
                db_data["education_list"] = parsed["education"]
            if parsed.get("work_experience"):
                db_data["work_experience_list"] = parsed["work_experience"]
            if parsed.get("projects"):
                db_data["project_list"] = parsed["projects"]
            if parsed.get("skills"):
                db_data["skill_list"] = parsed["skills"]
                extracted_skills = [s.get("name", "") for s in parsed["skills"] if s.get("name")]
            if parsed.get("self_evaluation"):
                db_data["self_evaluation"] = str(parsed["self_evaluation"])
        else:
            # LLM 失败 → 正则降级提取基本信息
            regex_result = regex_extract(raw_text)
            db_data.update(regex_result)
            db_data["self_evaluation"] = raw_text[:2000]

        # Step 5: 创建简历记录
        resume = await self.repo.create(user_id, db_data)
        await self.db.commit()

        return {
            "resume": self._to_dict(resume),
            "extracted_skills": extracted_skills,
            "parse_accuracy": parse_accuracy,
            "raw_text": raw_text,
        }

    def _to_dict(self, r) -> dict:
        """将 Resume 模型转为 API 返回格式"""
        return {
            "id": r.id,
            "name": r.name,
            "target_position": r.target_position,
            "personal_info": {
                "name": r.personal_name, "email": r.personal_email,
                "phone": r.personal_phone, "location": r.personal_location,
            },
            "job_intent": {
                "desired_position": r.desired_position or "",
                "desired_city": r.desired_city or "",
                "salary_expectation": r.salary_expectation or "",
                "work_mode": r.work_mode or "fulltime",
            },
            "education": r.education_list or [],
            "work_experience": r.work_experience_list or [],
            "projects": r.project_list or [],
            "skills": r.skill_list or [],
            "self_evaluation": r.self_evaluation or "",
            "source_file": r.source_file,
            "source_file_path": r.source_file_path,
            "raw_text": r.raw_text,
            "created_at": str(r.created_at) if r.created_at else None,
            "updated_at": str(r.updated_at) if r.updated_at else None,
        }
