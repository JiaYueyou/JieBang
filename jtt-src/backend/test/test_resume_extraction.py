from io import BytesIO

from docx import Document

from app.services.resume_service import (
    calculate_parse_completeness,
    extract_text_docx,
    regex_extract,
)


def test_extract_text_docx_reads_paragraph_table_header_and_footer():
    document = Document()
    document.add_paragraph("Python backend engineer")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "FastAPI"
    document.sections[0].header.paragraphs[0].text = "candidate@example.com"
    document.sections[0].footer.paragraphs[0].text = "13800138000"
    stream = BytesIO()
    document.save(stream)

    text = extract_text_docx(stream.getvalue())

    assert "Python backend engineer" in text
    assert "Skills | FastAPI" in text
    assert "candidate@example.com" in text
    assert "13800138000" in text


def test_regex_extract_contact_details():
    result = regex_extract("Test Candidate\ncandidate@example.com\n13800138000")

    assert result["personal_email"] == "candidate@example.com"
    assert result["personal_phone"] == "13800138000"


def test_parse_completeness_for_structured_resume():
    parsed = {
        "personal_info": {"name": "Test", "email": "test@example.com", "location": "Beijing"},
        "job_intent": {"desired_position": "Backend Engineer"},
        "education": [{"school": "Test University"}],
        "work_experience": [{"company": "Test Company"}],
        "projects": [{"name": "Test Project"}],
        "skills": [{"name": "Python"}],
    }

    assert calculate_parse_completeness(parsed, "resume text") == 1.0


def test_parse_completeness_is_zero_for_empty_text():
    assert calculate_parse_completeness(None, "  ") == 0.0
