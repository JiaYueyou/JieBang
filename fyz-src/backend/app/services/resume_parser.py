from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

from app.core.exceptions import InvalidParameterError


class ResumeParser:
    version = "resume-parser-v2"
    allowed_suffixes = {".txt", ".md", ".pdf", ".docx", ".png", ".jpg", ".jpeg"}
    image_suffixes = {".png", ".jpg", ".jpeg"}
    _ocr_engine: Any = None

    def parse(self, content: bytes, filename: str) -> tuple[str, list[str]]:
        if not content:
            raise InvalidParameterError("简历文件为空")
        if len(content) > 20 * 1024 * 1024:
            raise InvalidParameterError("简历文件不能超过 20MB")
        suffix = Path(filename).suffix.lower()
        if suffix not in self.allowed_suffixes:
            raise InvalidParameterError("仅支持 TXT、Markdown、PDF、DOCX、PNG、JPG 和 JPEG 简历")
        warnings: list[str] = []
        if suffix in {".txt", ".md"}:
            text = self._decode(content)
        elif suffix == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError as exc:
                raise InvalidParameterError("服务端未安装 PDF 解析依赖 pypdf") from exc
            text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages)
            if len(text.strip()) < 30:
                text = self._ocr_pdf(content)
                warnings.append("PDF 无有效文本层，已使用 OCR 识别。")
        elif suffix == ".docx":
            try:
                from docx import Document
            except ImportError as exc:
                raise InvalidParameterError("服务端未安装 Word 解析依赖 python-docx") from exc
            text = "\n".join(p.text for p in Document(io.BytesIO(content)).paragraphs)
        else:
            text = self._ocr_image(content)
            warnings.append("图片简历已使用 OCR 识别，请重点复核姓名、日期和数字。")
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())[:20000]
        if not text:
            raise InvalidParameterError("未能从简历中解析出文本")
        return text, warnings

    @classmethod
    def _rapid_ocr(cls):
        if cls._ocr_engine is None:
            try:
                from rapidocr_onnxruntime import RapidOCR
            except ImportError as exc:
                raise InvalidParameterError(
                    "服务端未安装图片 OCR 依赖 rapidocr-onnxruntime"
                ) from exc
            cls._ocr_engine = RapidOCR()
        return cls._ocr_engine

    def _ocr_image(self, content: bytes) -> str:
        text, _ = self.ocr_image_with_diagnostics(content)
        return text

    def ocr_image_with_diagnostics(
        self, content: bytes, *, optimized: bool = True
    ) -> tuple[str, dict[str, Any]]:
        """OCR an image and expose non-sensitive quality diagnostics.

        The optimized path compares the original image with an adaptive
        high-resolution/contrast-enhanced variant. It keeps the candidate
        with the strongest character-weighted OCR confidence and useful-text
        coverage instead of concatenating competing OCR outputs.
        """
        try:
            import numpy as np
            from PIL import Image, ImageEnhance, ImageOps
        except ImportError as exc:
            raise InvalidParameterError("服务端未安装图片解析依赖 Pillow/numpy") from exc
        try:
            image = ImageOps.exif_transpose(Image.open(io.BytesIO(content))).convert("RGB")
        except Exception as exc:
            raise InvalidParameterError("图片文件损坏或格式无法识别") from exc
        candidates: list[tuple[str, Any]] = [("original", image)]
        if optimized:
            scale = min(3.0, max(1.0, 1800 / max(1, image.width)))
            enhanced = image
            if scale > 1.05:
                enhanced = enhanced.resize(
                    (round(image.width * scale), round(image.height * scale)),
                    Image.Resampling.LANCZOS,
                )
            enhanced = ImageOps.autocontrast(enhanced, cutoff=1)
            enhanced = ImageEnhance.Sharpness(enhanced).enhance(1.35)
            candidates.append(("enhanced", enhanced))

        evaluated: list[dict[str, Any]] = []
        engine = self._rapid_ocr()
        for variant, candidate in candidates:
            result, _ = engine(np.asarray(candidate))
            rows = [
                row for row in (result or [])
                if len(row) >= 3 and str(row[1]).strip()
            ]
            text = self._normalize_ocr_text(
                "\n".join(str(row[1]).strip() for row in rows)
            )
            char_count = len(re.sub(r"\s+", "", text))
            weighted_confidence = (
                sum(float(row[2]) * len(str(row[1]).strip()) for row in rows)
                / max(1, sum(len(str(row[1]).strip()) for row in rows))
            )
            useful_ratio = (
                len(re.findall(r"[A-Za-z0-9\u3400-\u9fff]", text))
                / max(1, len(text))
            )
            selection_score = (
                weighted_confidence * 0.82
                + min(1.0, char_count / 500) * 0.12
                + useful_ratio * 0.06
            )
            evaluated.append({
                "variant": variant,
                "text": text,
                "line_count": len(rows),
                "character_count": char_count,
                "mean_confidence": round(weighted_confidence, 6),
                "selection_score": round(selection_score, 6),
                "width": candidate.width,
                "height": candidate.height,
            })
            if (
                optimized
                and variant == "original"
                and weighted_confidence >= 0.96
                and char_count >= 350
            ):
                # High-quality originals do not benefit enough from a second
                # inference pass to justify doubling upload latency.
                break
        selected = max(
            evaluated,
            key=lambda item: (
                item["selection_score"], item["character_count"],
            ),
        )
        diagnostics = {
            "selected_variant": selected["variant"],
            "line_count": selected["line_count"],
            "character_count": selected["character_count"],
            "mean_confidence": selected["mean_confidence"],
            "selection_score": selected["selection_score"],
            "candidates": [
                {key: value for key, value in item.items() if key != "text"}
                for item in evaluated
            ],
        }
        return selected["text"], diagnostics

    def _ocr_pdf(self, content: bytes) -> str:
        try:
            import pypdfium2 as pdfium
        except ImportError as exc:
            raise InvalidParameterError(
                "扫描 PDF 需要 pypdfium2 与 rapidocr-onnxruntime"
            ) from exc
        try:
            document = pdfium.PdfDocument(content)
            pages = [page.render(scale=2).to_pil() for page in document]
        except Exception as exc:
            raise InvalidParameterError("PDF 文件损坏或无法渲染") from exc
        try:
            import numpy as np
        except ImportError as exc:
            raise InvalidParameterError("服务端未安装 OCR 依赖 numpy") from exc
        engine = self._rapid_ocr()
        lines: list[str] = []
        for image in pages:
            result, _ = engine(np.asarray(image.convert("RGB")))
            lines.extend(row[1] for row in (result or []) if len(row) >= 2 and row[1].strip())
        return self._normalize_ocr_text("\n".join(lines))

    @staticmethod
    def _normalize_ocr_text(text: str) -> str:
        """Repair high-confidence OCR confusions in known technical tokens."""
        replacements = {
            "FastAPl": "FastAPI",
            "GraphQl": "GraphQL",
            "PostgreSQl": "PostgreSQL",
            "MySQl": "MySQL",
        }
        for source, target in replacements.items():
            text = text.replace(source, target)
        return text

    @staticmethod
    def _decode(content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                pass
        raise InvalidParameterError("无法识别文本文件编码")
